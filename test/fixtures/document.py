import io

import docx
import pytest
from app.models import Document, Author
from app.services.style_profile_service import StyleProfileService
from test.utils import TestingSessionLocal
from .authors import test_author
import hashlib

@pytest.fixture
def test_document(test_author: Author):
    example_text = "This is a sample document"

    db = TestingSessionLocal()
    author = db.get(Author, test_author.id)
    document = Document(
        authors=[author],
        filename="sample.docx",
        text=example_text,
        word_count=len(example_text.split()),
        content_hash=hashlib.sha256(example_text.encode("utf-8")).hexdigest(),
        **StyleProfileService().compute_document_stats(example_text),
    )

    db.add(document)
    db.commit()

    yield document


def _build_docx_bytes(paragraphs: list[str]) -> bytes:
    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def make_docx_upload():
    """Builds a real, in-memory .docx file usable as a multipart upload."""
    def _make(paragraphs=None, filename="sample.docx"):
        paragraphs = paragraphs if paragraphs is not None else ["Hello world, this is a sample document."]
        content = _build_docx_bytes(paragraphs)
        return filename, io.BytesIO(content), content
    return _make