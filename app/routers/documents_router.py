from fastapi import APIRouter, BackgroundTasks, UploadFile, Depends, HTTPException, Path, Query, status
from app.extensions import get_db
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from typing import Annotated
from app.models.author import Author
from app.models.document import Document
from app.schemas.document_response import DocumentResponse
from app.services.style_profile_service import StyleProfileService
from app.services.style_profile_rebuild_service import StyleProfileRebuildService
import docx
from docx.opc.exceptions import PackageNotFoundError
import hashlib
import io
import zipfile

router = APIRouter(
    prefix="/documents",
    tags=["Documents"]
)


@router.get("/{document_name}", response_model=DocumentResponse, status_code=status.HTTP_200_OK)
async def get_document(db: Annotated[Session, Depends(get_db)], document_name: str = Path(min_length=1)):
    document = db.query(Document).filter(Document.filename == document_name).first()

    if document is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found.")

    return document


def _get_or_create_author(db: Session, name: str) -> Author:
    title_cased_name = name.title()
    try:
        author = Author(name=title_cased_name)
        db.add(author)
        db.commit()
    except IntegrityError:
        db.rollback()
        author = db.query(Author).filter(Author.name == title_cased_name).first()
    return author


@router.post("/upload-known", response_model=DocumentResponse)
async def upload_document_known_author(
        db: Annotated[Session, Depends(get_db)],
        background_tasks: BackgroundTasks,
        file: UploadFile,
        author_names: Annotated[list[str], Query()]
):
    if not author_names:
        raise HTTPException(status_code=422, detail="At least one author_name is required")

    authors = [_get_or_create_author(db, name) for name in author_names]

    try:
        file_bytes = await file.read()
        doc = docx.Document(io.BytesIO(file_bytes))
        doc_text = " ".join([para.text for para in doc.paragraphs])
        word_count = sum(len(para.text.split()) for para in doc.paragraphs)
        content_hash = hashlib.sha256(doc_text.encode("utf-8")).hexdigest()
    except (PackageNotFoundError, zipfile.BadZipFile):
        raise HTTPException(status_code=422, detail="Trouble reading document. Not .docx")

    existing_document = db.query(Document).filter(Document.content_hash == content_hash).first()

    if existing_document is not None:
        new_authors = [author for author in authors if author not in existing_document.authors]
        if not new_authors:
            raise HTTPException(status_code=422, detail="Document already exists for these authors")
        existing_document.authors.extend(new_authors)
        db.commit()
        db.refresh(existing_document)
        background_tasks.add_task(StyleProfileRebuildService().rebuild_all_in_background)
        return existing_document

    document_stats = StyleProfileService().compute_document_stats(doc_text)

    try:
        new_document = Document(
            filename=file.filename,
            text=doc_text,
            word_count=word_count,
            content_hash=content_hash,
            authors=authors,
            **document_stats,
        )

        db.add(new_document)
        db.commit()
        db.refresh(new_document)
        background_tasks.add_task(StyleProfileRebuildService().rebuild_all_in_background)
        return new_document
    except IntegrityError:
        db.rollback()
        raise HTTPException(status_code=422, detail="Document already exists for these authors")



